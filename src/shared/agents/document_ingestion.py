"""Document Ingestion & Synthesis Agent

This module handles the parsing and synthesis of various document formats
into a unified ProjectBrief.md file.
"""


import sys
from pathlib import Path
# src 디렉토리를 Python path에 추가
src_dir = Path(__file__).parent.parent if 'src' in str(Path(__file__).parent) else Path(__file__).parent
while src_dir.name != 'src' and src_dir.parent != src_dir:
    src_dir = src_dir.parent
if src_dir.name == 'src':
    sys.path.insert(0, str(src_dir))

import os
import logging
from pathlib import Path
<<<<<<< HEAD
from typing import List, Dict, Any, Optional, Tuple
import re
=======
from typing import List, Dict, Any, Optional
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271
from datetime import datetime

import pdfplumber
from docx import Document
from PIL import Image
import google.generativeai as genai
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn

from shared.utils.config import Config
from shared.utils.file_utils import ensure_directory, read_text_file

logger = logging.getLogger(__name__)
console = Console()


class DocumentIngestionAgent:
    """Agent responsible for parsing and synthesizing project documentation."""
    
    def __init__(self, config: Config):
        self.config = config
        self.gemini_model = None
        self._setup_gemini()
        
    def _setup_gemini(self):
        """Initialize Gemini model for image analysis."""
        genai.configure(api_key=self.config.gemini_api_key)
        self.gemini_model = genai.GenerativeModel(self.config.gemini_model)
    
    def process_documents(self, docs_path: str = "docs") -> str:
        """Process all documents in the docs folder and create ProjectBrief.md
        
        Args:
            docs_path: Path to the documents folder
            
        Returns:
            Path to the generated ProjectBrief.md file
        """
        docs_dir = Path(docs_path)
        if not docs_dir.exists():
            raise FileNotFoundError(f"Documents directory not found: {docs_path}")
        
        console.print(f"[blue]Processing documents in {docs_path}...[/blue]")
        
        # Collect all documents
        documents = self._collect_documents(docs_dir)
        
        # Process each document type
        all_content = []
        
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=console,
        ) as progress:
            
            # Process text files
            task = progress.add_task("Processing text files...", total=len(documents['text']))
            for file_path in documents['text']:
                content = self._process_text_file(file_path)
                if content:
                    all_content.append(self._format_section(file_path.name, content))
                progress.advance(task)
            
            # Process PDFs
            task = progress.add_task("Processing PDF files...", total=len(documents['pdf']))
            for file_path in documents['pdf']:
                content = self._process_pdf_file(file_path)
                if content:
                    all_content.append(self._format_section(file_path.name, content))
                progress.advance(task)
            
            # Process Word documents
            task = progress.add_task("Processing Word documents...", total=len(documents['docx']))
            for file_path in documents['docx']:
                content = self._process_docx_file(file_path)
                if content:
                    all_content.append(self._format_section(file_path.name, content))
                progress.advance(task)
            
            # Process images
            task = progress.add_task("Processing images...", total=len(documents['images']))
            for file_path in documents['images']:
                content = self._process_image_file(file_path)
                if content:
                    all_content.append(self._format_section(f"{file_path.name} (Image Analysis)", content))
                progress.advance(task)
        
        # Generate ProjectBrief.md
        project_brief_path = self._generate_project_brief(all_content)
        
        console.print(f"[green]✓ ProjectBrief.md generated at: {project_brief_path}[/green]")
        return project_brief_path
    
    def _collect_documents(self, docs_dir: Path) -> Dict[str, List[Path]]:
        """Collect and categorize all documents in the directory."""
        documents = {
            'text': [],
            'pdf': [],
            'docx': [],
            'images': []
        }
        
        for file_path in docs_dir.rglob('*'):
            if file_path.is_file():
                suffix = file_path.suffix.lower()
                
                if suffix in ['.md', '.txt', '.rst']:
                    documents['text'].append(file_path)
                elif suffix == '.pdf':
                    documents['pdf'].append(file_path)
                elif suffix in ['.docx', '.doc']:
                    documents['docx'].append(file_path)
                elif suffix in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                    documents['images'].append(file_path)
        
        return documents
    
    def _process_text_file(self, file_path: Path) -> Optional[str]:
        """Process text-based files (Markdown, plain text, etc.)."""
        try:
            return read_text_file(str(file_path))
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            return None
    
    def _process_pdf_file(self, file_path: Path) -> Optional[str]:
<<<<<<< HEAD
        """Extract text, tables, and structure from PDF files with enhanced processing."""
        try:
            content = []
            document_structure = {"sections": [], "tables": [], "images": []}
            
            with pdfplumber.open(file_path) as pdf:
                # Extract document metadata
                metadata = pdf.metadata
                if metadata:
                    content.append(f"# Document: {metadata.get('Title', file_path.name)}")
                    if metadata.get('Author'):
                        content.append(f"Author: {metadata.get('Author')}")
                    content.append("\n")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Enhanced text extraction with layout preservation
                    page_content = self._extract_page_with_structure(page, page_num)
                    if page_content:
                        content.append(page_content)
                    
                    # Enhanced table extraction
                    tables = self._extract_enhanced_tables(page, page_num)
                    if tables:
                        content.extend(tables)
                        document_structure["tables"].extend(tables)
                    
                    # Extract images and their context
                    images = self._extract_page_images(page, page_num)
                    if images:
                        document_structure["images"].extend(images)
            
            # Add document structure summary
            if document_structure["tables"]:
                content.append(f"\n## Document contains {len(document_structure['tables'])} tables")
            if document_structure["images"]:
                content.append(f"## Document contains {len(document_structure['images'])} images/diagrams")
=======
        """Extract text and tables from PDF files."""
        try:
            content = []
            
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text()
                    if text:
                        content.append(f"Page {page_num}:\n{text}")
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_md = self._table_to_markdown(table)
                            content.append(f"\nTable {table_idx + 1} on Page {page_num}:\n{table_md}")
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271
            
            return "\n\n".join(content)
        
        except Exception as e:
            logger.error(f"Error processing PDF file {file_path}: {e}")
            return None
    
    def _process_docx_file(self, file_path: Path) -> Optional[str]:
        """Extract text and tables from Word documents."""
        try:
            doc = Document(file_path)
            content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    table_md = self._table_to_markdown(table_data)
                    content.append(f"\nTable {table_idx + 1}:\n{table_md}")
            
            return "\n\n".join(content)
        
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {e}")
            return None
    
    def _process_image_file(self, file_path: Path) -> Optional[str]:
<<<<<<< HEAD
        """Analyze images using Gemini's multimodal capabilities with enhanced UI component extraction."""
        try:
            image = Image.open(file_path)
            
            # Determine image type based on filename and content
            image_type = self._determine_image_type(file_path, image)
            
            # Use specialized prompts based on image type
            if image_type == "ui_mockup":
                prompt = self._get_ui_analysis_prompt()
            elif image_type == "diagram":
                prompt = self._get_diagram_analysis_prompt()
            elif image_type == "screenshot":
                prompt = self._get_screenshot_analysis_prompt()
            else:
                prompt = self._get_general_image_prompt()
            
            response = self.gemini_model.generate_content([prompt, image])
            
            # Post-process the response to extract structured information
            structured_content = self._structure_image_analysis(response.text, image_type)
            
            return f"### Image Analysis: {file_path.name}\n\n{structured_content}"
=======
        """Analyze images using Gemini's multimodal capabilities."""
        try:
            image = Image.open(file_path)
            
            prompt = """Analyze this image and provide a detailed description. 
            If this is a UI mockup, wireframe, or design document, describe:
            - The layout and structure
            - UI components present (buttons, forms, navigation, etc.)
            - Any text or labels visible
            - The apparent functionality or purpose
            - Color scheme and visual style if relevant
            
            If this is a diagram or flowchart, describe:
            - The type of diagram
            - Components and their relationships
            - Any process flow or data flow shown
            - Key information conveyed
            
            Be as specific and detailed as possible."""
            
            response = self.gemini_model.generate_content([prompt, image])
            return response.text
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271
        
        except Exception as e:
            logger.error(f"Error processing image file {file_path}: {e}")
            return None
    
<<<<<<< HEAD
    def _determine_image_type(self, file_path: Path, image: Image) -> str:
        """Determine the type of image based on filename and characteristics."""
        filename_lower = file_path.name.lower()
        
        # Check filename patterns
        if any(keyword in filename_lower for keyword in ["mockup", "wireframe", "ui", "design", "screen"]):
            return "ui_mockup"
        elif any(keyword in filename_lower for keyword in ["diagram", "flow", "chart", "architecture"]):
            return "diagram"
        elif any(keyword in filename_lower for keyword in ["screenshot", "capture"]):
            return "screenshot"
        
        # Could add image analysis here (e.g., aspect ratio, color distribution)
        # For now, default to general
        return "general"
    
    def _get_ui_analysis_prompt(self) -> str:
        """Get specialized prompt for UI mockup analysis."""
        return """Analyze this UI mockup/wireframe and extract the following information:

1. **Page/Screen Identification**:
   - What page or screen is this? (e.g., login, dashboard, profile)
   - Primary purpose and functionality

2. **Layout Structure**:
   - Overall layout pattern (header-body-footer, sidebar layout, grid, etc.)
   - Number of main sections/regions
   - Responsive design considerations visible

3. **UI Components** (list each with details):
   - Navigation elements (menus, breadcrumbs, tabs)
   - Forms (input fields, labels, validation hints)
   - Buttons (text, style, apparent actions)
   - Data display (tables, lists, cards)
   - Media elements (images, videos, icons)
   - Modals/dialogs/popups

4. **Visual Design**:
   - Color scheme (primary, secondary, accent colors)
   - Typography hierarchy visible
   - Spacing and alignment patterns
   - Any dark/light theme indicators

5. **User Interactions**:
   - Clickable elements and their likely actions
   - Form workflows
   - Navigation flows
   - State changes (hover, active, disabled states)

6. **Content Structure**:
   - Headings and their hierarchy
   - Text content organization
   - Data presentation format

7. **Technical Implementation Hints**:
   - Suggested component names
   - State management needs
   - API endpoints that might be needed
   - Responsive breakpoints

Format the response with clear sections and bullet points for easy parsing."""
    
    def _get_diagram_analysis_prompt(self) -> str:
        """Get specialized prompt for diagram analysis."""
        return """Analyze this diagram and provide:

1. **Diagram Type**: (flowchart, sequence diagram, ER diagram, architecture diagram, etc.)

2. **Components/Entities**:
   - List all boxes, nodes, or entities
   - Their labels and descriptions
   - Any properties or attributes shown

3. **Relationships/Connections**:
   - All connections between components
   - Direction of flow (if applicable)
   - Labels on connections
   - Type of relationships (one-to-many, etc.)

4. **Process Flow** (if applicable):
   - Starting point
   - Decision points
   - End points
   - Alternative paths

5. **Technical Details**:
   - Any technical specifications mentioned
   - Data types, protocols, or standards
   - System boundaries

6. **Key Insights**:
   - Main purpose of the diagram
   - Critical paths or components
   - Dependencies highlighted

Format as structured text with clear sections."""
    
    def _get_screenshot_analysis_prompt(self) -> str:
        """Get specialized prompt for screenshot analysis."""
        return """Analyze this screenshot and identify:

1. **Application/Website Context**:
   - What application or website is this?
   - Current page/screen

2. **Functional Elements**:
   - All interactive elements visible
   - Current state of the application
   - Any data being displayed

3. **UI Patterns**:
   - Design patterns used
   - Component library indicators
   - Styling approach

4. **Technical Observations**:
   - Framework indicators (if any)
   - Responsive design elements
   - Accessibility features visible

Format findings in a structured way."""
    
    def _get_general_image_prompt(self) -> str:
        """Get general image analysis prompt."""
        return """Analyze this image and provide a comprehensive description including:

1. **Content Description**: What is shown in the image
2. **Purpose**: What this image is meant to convey or document
3. **Technical Details**: Any technical information visible
4. **Relationships**: How this might relate to a software project
5. **Implementation Hints**: Any insights for implementation

Be thorough and specific in your analysis."""
    
    def _structure_image_analysis(self, analysis_text: str, image_type: str) -> str:
        """Structure the image analysis response for better parsing."""
        if not analysis_text:
            return "No analysis available"
        
        # Add image type header
        structured = f"**Image Type**: {image_type.replace('_', ' ').title()}\n\n"
        structured += analysis_text
        
        # Add extraction markers for key components
        if image_type == "ui_mockup":
            # Try to extract component list
            components = self._extract_ui_components(analysis_text)
            if components:
                structured += "\n\n**Extracted Components for Implementation**:\n"
                for comp in components:
                    structured += f"- {comp}\n"
        
        return structured
    
    def _extract_ui_components(self, analysis_text: str) -> List[str]:
        """Extract UI component names from analysis text."""
        components = []
        
        # Look for component mentions
        component_patterns = [
            r"(?:button|btn)(?:\s+(?:labeled?|text|with))?[:\s]+[\"']?([^\"'\n]+)[\"']?",
            r"(?:input|field|textbox)(?:\s+(?:for|labeled?))?\s*[:\s]+([^\n]+)",
            r"(?:navigation|nav|menu)\s*[:\s]+([^\n]+)",
            r"(?:table|list|grid)\s+(?:showing|displaying|for)\s+([^\n]+)",
        ]
        
        for pattern in component_patterns:
            matches = re.findall(pattern, analysis_text, re.IGNORECASE)
            components.extend(matches)
        
        # Clean and deduplicate
        components = list(set(comp.strip() for comp in components if comp.strip()))
        
        return components
    
    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """Convert table data to Markdown format with enhanced formatting."""
        if not table_data or not table_data[0]:
            return ""
        
        # Clean and normalize table data
        cleaned_data = []
        for row in table_data:
            cleaned_row = [str(cell).strip() if cell else "" for cell in row]
            if any(cleaned_row):  # Skip empty rows
                cleaned_data.append(cleaned_row)
        
        if not cleaned_data:
            return ""
        
        # Ensure all rows have the same number of columns
        max_cols = max(len(row) for row in cleaned_data)
        normalized_data = [row + [""] * (max_cols - len(row)) for row in cleaned_data]
        
        # Create header
        header = "| " + " | ".join(normalized_data[0]) + " |"
        separator = "| " + " | ".join("---" for _ in range(max_cols)) + " |"
        
        # Create rows
        rows = []
        for row in normalized_data[1:]:
            row_str = "| " + " | ".join(row) + " |"
=======
    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """Convert table data to Markdown format."""
        if not table_data or not table_data[0]:
            return ""
        
        # Create header
        header = "| " + " | ".join(str(cell) for cell in table_data[0]) + " |"
        separator = "| " + " | ".join("---" for _ in table_data[0]) + " |"
        
        # Create rows
        rows = []
        for row in table_data[1:]:
            row_str = "| " + " | ".join(str(cell) for cell in row) + " |"
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271
            rows.append(row_str)
        
        return "\n".join([header, separator] + rows)
    
<<<<<<< HEAD
    def _extract_page_with_structure(self, page, page_num: int) -> Optional[str]:
        """Extract text from page with structure and formatting preservation."""
        try:
            # Extract text with layout information
            text_with_layout = page.extract_text_simple(x_tolerance=2, y_tolerance=3)
            if not text_with_layout:
                return None
            
            # Extract text with character-level information for heading detection
            chars = page.chars
            if chars:
                # Group characters by similar formatting (font, size)
                formatted_sections = self._group_by_formatting(chars)
                
                # Identify headings based on font size and style
                headings = self._identify_headings(formatted_sections)
                
                # Reconstruct text with proper heading markers
                structured_text = self._reconstruct_with_headings(text_with_layout, headings)
                return f"## Page {page_num}\n\n{structured_text}"
            else:
                return f"## Page {page_num}\n\n{text_with_layout}"
                
        except Exception as e:
            logger.debug(f"Could not extract structured text from page {page_num}: {e}")
            # Fallback to simple extraction
            text = page.extract_text()
            return f"## Page {page_num}\n\n{text}" if text else None
    
    def _extract_enhanced_tables(self, page, page_num: int) -> List[str]:
        """Extract tables with context and improved formatting."""
        tables_content = []
        try:
            # Configure table extraction settings for better accuracy
            table_settings = {
                "vertical_strategy": "lines",
                "horizontal_strategy": "lines",
                "explicit_vertical_lines": [],
                "explicit_horizontal_lines": [],
                "snap_tolerance": 3,
                "join_tolerance": 3,
                "edge_min_length": 3,
                "min_words_vertical": 1,
                "min_words_horizontal": 1,
                "text_tolerance": 3,
            }
            
            tables = page.extract_tables(table_settings)
            
            for table_idx, table in enumerate(tables):
                if table and self._is_valid_table(table):
                    # Look for table caption/title above the table
                    table_title = self._find_table_caption(page, table_idx)
                    
                    table_md = self._table_to_markdown(table)
                    if table_title:
                        tables_content.append(f"\n### {table_title}\n{table_md}")
                    else:
                        tables_content.append(f"\n### Table {table_idx + 1} (Page {page_num})\n{table_md}")
                        
        except Exception as e:
            logger.debug(f"Error extracting tables from page {page_num}: {e}")
            
        return tables_content
    
    def _extract_page_images(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Extract information about images on the page."""
        images_info = []
        try:
            # Note: PDFPlumber doesn't directly extract images, but we can identify image regions
            # This is a placeholder for image region detection
            # In a full implementation, you might use other libraries like PyMuPDF for actual image extraction
            
            # For now, we'll look for figure references in the text
            text = page.extract_text()
            if text:
                figure_refs = re.findall(r'(?:Figure|Fig\\.?)\\s*(\\d+[\\w.-]*)', text, re.IGNORECASE)
                for fig_ref in figure_refs:
                    images_info.append({
                        "page": page_num,
                        "reference": f"Figure {fig_ref}",
                        "type": "figure_reference"
                    })
                    
        except Exception as e:
            logger.debug(f"Error extracting image info from page {page_num}: {e}")
            
        return images_info
    
    def _group_by_formatting(self, chars: List[Dict]) -> List[Dict[str, Any]]:
        """Group characters by similar formatting (font, size)."""
        if not chars:
            return []
            
        groups = []
        current_group = {"chars": [chars[0]], "fontname": chars[0].get("fontname"), "size": chars[0].get("size")}
        
        for char in chars[1:]:
            # Check if formatting matches current group
            if (char.get("fontname") == current_group["fontname"] and 
                abs(char.get("size", 0) - current_group["size"]) < 0.5):
                current_group["chars"].append(char)
            else:
                # Start new group
                groups.append(current_group)
                current_group = {"chars": [char], "fontname": char.get("fontname"), "size": char.get("size")}
                
        groups.append(current_group)
        return groups
    
    def _identify_headings(self, formatted_sections: List[Dict]) -> List[Dict[str, Any]]:
        """Identify headings based on font size and style."""
        if not formatted_sections:
            return []
            
        # Calculate average font size
        sizes = [section["size"] for section in formatted_sections if section.get("size")]
        if not sizes:
            return []
            
        avg_size = sum(sizes) / len(sizes)
        
        headings = []
        for section in formatted_sections:
            size = section.get("size", 0)
            # Consider text as heading if it's significantly larger than average
            if size > avg_size * 1.2:
                text = "".join(char.get("text", "") for char in section["chars"])
                if text.strip():
                    headings.append({
                        "text": text.strip(),
                        "size": size,
                        "level": self._determine_heading_level(size, avg_size)
                    })
                    
        return headings
    
    def _determine_heading_level(self, size: float, avg_size: float) -> int:
        """Determine heading level based on font size."""
        ratio = size / avg_size
        if ratio > 1.5:
            return 1  # H1
        elif ratio > 1.3:
            return 2  # H2
        elif ratio > 1.2:
            return 3  # H3
        else:
            return 4  # H4
    
    def _reconstruct_with_headings(self, text: str, headings: List[Dict]) -> str:
        """Reconstruct text with proper heading markers."""
        if not headings:
            return text
            
        result = text
        # Replace identified headings with markdown format
        for heading in headings:
            heading_text = heading["text"]
            level = heading["level"]
            markdown_heading = "#" * level + " " + heading_text
            # Replace the heading in text (being careful with partial matches)
            result = result.replace(heading_text, markdown_heading)
            
        return result
    
    def _is_valid_table(self, table: List[List]) -> bool:
        """Check if extracted table is valid and not just random text."""
        if not table or len(table) < 2:
            return False
            
        # Check if table has consistent column count
        col_counts = [len(row) for row in table if row]
        if not col_counts:
            return False
            
        # Most rows should have similar column count
        most_common_count = max(set(col_counts), key=col_counts.count)
        consistency_ratio = col_counts.count(most_common_count) / len(col_counts)
        
        return consistency_ratio > 0.7
    
    def _find_table_caption(self, page, table_idx: int) -> Optional[str]:
        """Try to find table caption or title."""
        try:
            text = page.extract_text()
            if not text:
                return None
                
            # Look for common table caption patterns
            patterns = [
                f"Table\\s*{table_idx + 1}[\\s:.-]*(.*?)(?:\\n|$)",
                f"TABLE\\s*{table_idx + 1}[\\s:.-]*(.*?)(?:\\n|$)",
                f"Table\\s*\\w+[\\s:.-]*(.*?)(?:\\n|$)"
            ]
            
            for pattern in patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match and match.group(1):
                    return f"Table {table_idx + 1}: {match.group(1).strip()}"
                    
            return None
            
        except Exception:
            return None
    
=======
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271
    def _format_section(self, title: str, content: str) -> str:
        """Format a content section for the ProjectBrief."""
        return f"""
## Source: {title}

{content}

---
"""
    
    def _generate_project_brief(self, all_content: List[str]) -> str:
<<<<<<< HEAD
        """Generate the final ProjectBrief.md file with semantic analysis."""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        # Perform semantic analysis on all content
        analysis_results = self._analyze_document_semantics(all_content)
        
        # Create structured sections based on analysis
        structured_content = self._create_structured_brief(all_content, analysis_results)
        
=======
        """Generate the final ProjectBrief.md file."""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271
        brief_content = f"""# Project Brief
Generated on: {timestamp}

This document is a synthesized compilation of all project documentation found in the docs folder.
It serves as the primary input for the AI planning and code generation process.

<<<<<<< HEAD
## Executive Summary
{analysis_results.get('executive_summary', 'No summary available')}

## Document Analysis

### Project Type
{analysis_results.get('project_type', 'Unknown')}

### Key Requirements
{self._format_requirements(analysis_results.get('requirements', []))}

### Technical Constraints
{self._format_constraints(analysis_results.get('constraints', []))}

### Identified Technologies
{self._format_technologies(analysis_results.get('technologies', []))}

---

## Detailed Documentation

{structured_content}
=======
---

{"".join(all_content)}
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271

## End of Project Brief

This concludes the synthesized project documentation. The Master Planner Agent will use this 
information to design the complete software architecture and create detailed implementation tasks.
"""
        
        # Save to file
        output_path = Path("ProjectBrief.md")
        output_path.write_text(brief_content, encoding='utf-8')
        
<<<<<<< HEAD
        return str(output_path)
    
    def _analyze_document_semantics(self, all_content: List[str]) -> Dict[str, Any]:
        """Analyze document content to extract semantic information."""
        combined_text = "\n".join(all_content)
        
        analysis = {
            "executive_summary": self._generate_executive_summary(combined_text),
            "project_type": self._identify_project_type(combined_text),
            "requirements": self._extract_requirements(combined_text),
            "constraints": self._extract_constraints(combined_text),
            "technologies": self._extract_technologies(combined_text),
            "key_features": self._extract_key_features(combined_text),
            "user_stories": self._extract_user_stories(combined_text)
        }
        
        return analysis
    
    def _generate_executive_summary(self, text: str) -> str:
        """Generate an executive summary using Gemini."""
        try:
            prompt = """Based on the following project documentation, generate a concise executive summary 
            (3-5 sentences) that captures the essence of the project, its main goals, and target audience:
            
            {text[:3000]}
            """
            
            response = self.gemini_model.generate_content(prompt.format(text=text))
            return response.text
        except Exception as e:
            logger.error(f"Error generating executive summary: {e}")
            return "Unable to generate executive summary."
    
    def _identify_project_type(self, text: str) -> str:
        """Identify the type of project from the documentation."""
        text_lower = text.lower()
        
        # Check for project type indicators
        project_types = {
            "web_app": ["web application", "webapp", "website", "frontend", "backend", "full-stack"],
            "mobile_app": ["mobile app", "ios", "android", "react native", "flutter"],
            "api": ["api", "rest", "graphql", "microservice", "endpoint"],
            "desktop_app": ["desktop", "electron", "windows app", "mac app"],
            "library": ["library", "package", "module", "framework"],
            "data_system": ["data pipeline", "etl", "analytics", "dashboard", "reporting"],
            "ml_system": ["machine learning", "ml", "ai", "neural network", "model"],
            "iot": ["iot", "embedded", "sensor", "device"],
            "blockchain": ["blockchain", "smart contract", "dapp", "crypto"]
        }
        
        scores = {}
        for ptype, keywords in project_types.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            if score > 0:
                scores[ptype] = score
        
        if scores:
            return max(scores.items(), key=lambda x: x[1])[0].replace("_", " ").title()
        return "General Software Project"
    
    def _extract_requirements(self, text: str) -> List[str]:
        """Extract functional and non-functional requirements."""
        requirements = []
        
        # Look for requirement patterns
        patterns = [
            r"(?:must|should|shall|required to)\s+([^.!?]+[.!?])",
            r"(?:requirement[s]?|feature[s]?):\s*([^.!?]+[.!?])",
            r"- (?:The system|The application|The user)\s+(?:must|should|can)\s+([^.!?\n]+)",
            r"\d+\.\s*(?:The system|The application)\s+(?:must|should)\s+([^.!?\n]+)"
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            requirements.extend([match.strip() for match in matches])
        
        # Remove duplicates and clean
        requirements = list(set(req for req in requirements if len(req) > 20))
        
        return requirements[:15]  # Limit to top 15
    
    def _extract_constraints(self, text: str) -> List[str]:
        """Extract technical and business constraints."""
        constraints = []
        
        # Look for constraint patterns
        patterns = [
            r"(?:constraint[s]?|limitation[s]?|restriction[s]?):\s*([^.!?]+[.!?])",
            r"(?:must not|cannot|should not|limited to)\s+([^.!?]+[.!?])",
            r"(?:maximum|minimum|no more than|at least)\s+([^.!?]+[.!?])",
            r"(?:performance|security|scalability)\s+(?:requirement|constraint):\s*([^.!?]+[.!?])"
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            constraints.extend([match.strip() for match in matches])
        
        # Also look for technology constraints
        if "must use" in text.lower() or "required to use" in text.lower():
            tech_constraints = re.findall(r"(?:must use|required to use)\s+([^.!?]+[.!?])", text, re.IGNORECASE)
            constraints.extend([f"Technology constraint: {tc.strip()}" for tc in tech_constraints])
        
        return list(set(constraints))[:10]
    
    def _extract_technologies(self, text: str) -> List[str]:
        """Extract mentioned technologies, frameworks, and tools."""
        technologies = []
        
        # Common technology patterns
        tech_keywords = {
            # Frontend
            "React", "Angular", "Vue", "Next.js", "Nuxt", "Svelte", "TypeScript", "JavaScript",
            "Tailwind CSS", "Bootstrap", "Material-UI", "Chakra UI",
            # Backend
            "Node.js", "Express", "FastAPI", "Django", "Flask", "Spring Boot", "Ruby on Rails",
            "ASP.NET", "Laravel", "NestJS",
            # Databases
            "MySQL", "PostgreSQL", "MongoDB", "Redis", "SQLite", "Oracle", "SQL Server",
            "DynamoDB", "Cassandra", "Elasticsearch",
            # Cloud/DevOps
            "AWS", "Azure", "Google Cloud", "Docker", "Kubernetes", "Jenkins", "GitHub Actions",
            "Terraform", "Ansible",
            # Languages
            "Python", "Java", "C#", "Go", "Rust", "PHP", "Ruby", "Swift", "Kotlin",
            # Other
            "GraphQL", "REST API", "gRPC", "WebSocket", "JWT", "OAuth", "SAML"
        }
        
        # Find technologies in text
        for tech in tech_keywords:
            if tech.lower() in text.lower():
                technologies.append(tech)
        
        # Also look for explicit technology mentions
        tech_patterns = [
            r"(?:using|built with|powered by|based on)\s+([A-Za-z0-9\s\.\-]+?)(?:\s+and|\s*,|\s*\.)",
            r"(?:technology|framework|library):\s*([A-Za-z0-9\s\.\-]+?)(?:\s*,|\s*\.|\s*\n)"
        ]
        
        for pattern in tech_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            technologies.extend([match.strip() for match in matches if len(match.strip()) < 30])
        
        return list(set(technologies))
    
    def _extract_key_features(self, text: str) -> List[str]:
        """Extract key features of the system."""
        features = []
        
        # Look for feature patterns
        patterns = [
            r"(?:feature[s]?|functionality|capabilities):\s*([^.!?]+[.!?])",
            r"- ([A-Z][^.!?\n]+(?:feature|functionality|capability))",
            r"\d+\.\s*([A-Z][^.!?\n]+(?:management|system|module|component))"
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            features.extend([match.strip() for match in matches])
        
        return list(set(features))[:20]
    
    def _extract_user_stories(self, text: str) -> List[str]:
        """Extract user stories or use cases."""
        stories = []
        
        # Look for user story patterns
        patterns = [
            r"As a[n]?\s+([^,]+),\s*I (?:want|need|should be able to)\s+([^.!?]+)",
            r"User[s]?\s+(?:can|should be able to|must be able to)\s+([^.!?]+[.!?])",
            r"(?:Use case|User story):\s*([^.!?\n]+)"
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if isinstance(matches[0], tuple) if matches else False:
                # Handle tuple matches from the first pattern
                stories.extend([f"As a {match[0]}, I want to {match[1]}" for match in matches])
            else:
                stories.extend([match.strip() for match in matches])
        
        return list(set(stories))[:15]
    
    def _create_structured_brief(self, all_content: List[str], analysis: Dict[str, Any]) -> str:
        """Create a structured version of the content with proper organization."""
        # Group content by type
        grouped_content = {
            "requirements": [],
            "technical": [],
            "ui_ux": [],
            "business": [],
            "other": []
        }
        
        # Categorize each section
        for content in all_content:
            content_lower = content.lower()
            if any(keyword in content_lower for keyword in ["requirement", "feature", "user story", "use case"]):
                grouped_content["requirements"].append(content)
            elif any(keyword in content_lower for keyword in ["technical", "api", "database", "architecture"]):
                grouped_content["technical"].append(content)
            elif any(keyword in content_lower for keyword in ["ui", "ux", "design", "mockup", "wireframe"]):
                grouped_content["ui_ux"].append(content)
            elif any(keyword in content_lower for keyword in ["business", "market", "revenue", "customer"]):
                grouped_content["business"].append(content)
            else:
                grouped_content["other"].append(content)
        
        # Build structured content
        structured = []
        
        if grouped_content["requirements"]:
            structured.append("### Requirements Documentation\n")
            structured.extend(grouped_content["requirements"])
        
        if grouped_content["technical"]:
            structured.append("\n### Technical Documentation\n")
            structured.extend(grouped_content["technical"])
        
        if grouped_content["ui_ux"]:
            structured.append("\n### UI/UX Documentation\n")
            structured.extend(grouped_content["ui_ux"])
        
        if grouped_content["business"]:
            structured.append("\n### Business Documentation\n")
            structured.extend(grouped_content["business"])
        
        if grouped_content["other"]:
            structured.append("\n### Additional Documentation\n")
            structured.extend(grouped_content["other"])
        
        return "\n".join(structured)
    
    def _format_requirements(self, requirements: List[str]) -> str:
        """Format requirements list for the brief."""
        if not requirements:
            return "No specific requirements identified."
        
        formatted = []
        for i, req in enumerate(requirements, 1):
            formatted.append(f"{i}. {req}")
        
        return "\n".join(formatted)
    
    def _format_constraints(self, constraints: List[str]) -> str:
        """Format constraints list for the brief."""
        if not constraints:
            return "No specific constraints identified."
        
        formatted = []
        for i, constraint in enumerate(constraints, 1):
            formatted.append(f"{i}. {constraint}")
        
        return "\n".join(formatted)
    
    def _format_technologies(self, technologies: List[str]) -> str:
        """Format technologies list for the brief."""
        if not technologies:
            return "No specific technologies identified."
        
        # Group technologies by category
        categorized = {
            "Frontend": [],
            "Backend": [],
            "Database": [],
            "DevOps/Cloud": [],
            "Other": []
        }
        
        frontend_keywords = ["react", "angular", "vue", "next", "nuxt", "svelte", "tailwind", "bootstrap"]
        backend_keywords = ["node", "express", "django", "flask", "fastapi", "spring", "rails", "laravel"]
        database_keywords = ["mysql", "postgres", "mongodb", "redis", "sqlite", "oracle", "dynamodb"]
        devops_keywords = ["aws", "azure", "google cloud", "docker", "kubernetes", "jenkins", "terraform"]
        
        for tech in technologies:
            tech_lower = tech.lower()
            if any(keyword in tech_lower for keyword in frontend_keywords):
                categorized["Frontend"].append(tech)
            elif any(keyword in tech_lower for keyword in backend_keywords):
                categorized["Backend"].append(tech)
            elif any(keyword in tech_lower for keyword in database_keywords):
                categorized["Database"].append(tech)
            elif any(keyword in tech_lower for keyword in devops_keywords):
                categorized["DevOps/Cloud"].append(tech)
            else:
                categorized["Other"].append(tech)
        
        formatted = []
        for category, techs in categorized.items():
            if techs:
                formatted.append(f"**{category}**: {', '.join(techs)}")
        
        return "\n".join(formatted) if formatted else "No specific technologies identified."
=======
        return str(output_path)
>>>>>>> b6976b308e82b1aa019bf18e57915c15ddabb271
